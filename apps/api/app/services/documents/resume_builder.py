from __future__ import annotations

import io
import logging
import os

logger = logging.getLogger(__name__)


def build_docx(content_json: dict, full_name: str) -> bytes:
    """Return DOCX bytes from a tailored resume JSON."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    contact = content_json.get("contact") or {}
    name = contact.get("full_name") or full_name

    doc = Document()

    # --- Name header ---
    heading = doc.add_heading(name, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Contact line ---
    contact_parts = [p for p in [
        contact.get("email"), contact.get("phone"), contact.get("location"),
    ] if p]
    link_parts = [p for p in [
        contact.get("linkedin_url"), contact.get("github_url"), contact.get("portfolio_url"),
    ] if p]

    if contact_parts:
        cp = doc.add_paragraph(" | ".join(contact_parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(9)
    if link_parts:
        lp = doc.add_paragraph(" | ".join(link_parts))
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in lp.runs:
            run.font.size = Pt(9)

    if content_json.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(content_json["summary"])

    if content_json.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(content_json["skills"]))

    if content_json.get("experience"):
        doc.add_heading("Experience", level=1)
        for exp in content_json["experience"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            run.bold = True
            dates = f"{exp.get('start_date') or ''} – {exp.get('end_date') or 'Present'}"
            dp = doc.add_paragraph(dates)
            if dp.runs:
                dp.runs[0].italic = True
                dp.runs[0].font.size = Pt(9.5)
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    if content_json.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for cert in content_json["certifications"]:
            parts = [cert.get("name", "")]
            if cert.get("issuer"):
                parts.append(cert["issuer"])
            if cert.get("year"):
                parts.append(f"({cert['year']})")
            doc.add_paragraph(" — ".join(parts))

    if content_json.get("projects"):
        doc.add_heading("Projects", level=1)
        for proj in content_json["projects"]:
            p = doc.add_paragraph()
            p.add_run(proj.get("name", "")).bold = True
            if proj.get("technologies"):
                p.add_run(f" — {', '.join(proj['technologies'])}")
            if proj.get("description"):
                doc.add_paragraph(proj["description"])

    if content_json.get("education"):
        doc.add_heading("Education", level=1)
        for edu in content_json["education"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.get('degree', '')} — {edu.get('institution', '')}")
            run.bold = True
            if edu.get("graduation_year"):
                doc.add_paragraph(edu["graduation_year"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(content_json: dict, full_name: str) -> bytes:
    """Return PDF bytes rendered from the resume HTML template via WeasyPrint."""
    from jinja2 import Environment, FileSystemLoader
    from weasyprint import HTML

    templates_dir = os.path.join(os.path.dirname(__file__), "../../templates")
    env = Environment(loader=FileSystemLoader(os.path.abspath(templates_dir)))
    template = env.get_template("resume.html")
    # Pass full_name as fallback; contact info is also embedded in content_json["contact"]
    contact = content_json.get("contact") or {}
    html_str = template.render(
        full_name=contact.get("full_name") or full_name,
        resume=content_json,
    )

    pdf_bytes: bytes = HTML(string=html_str).write_pdf()
    return pdf_bytes
