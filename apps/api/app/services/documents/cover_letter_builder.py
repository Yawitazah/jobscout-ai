"""
Cover letter PDF + DOCX renderers.

Takes the body paragraphs the tailoring service produced and wraps them in a
standard cover letter layout:

    Lorenzo White                            (name, bold)
    Rock Hill, SC                            (contact lines)
    215-789-5593
    LorenzoR.White@gmail.com
    linkedin.com/in/lorenzowhite

    May 15, 2026                             (today's date)

    Hiring Team, Stripe                      (if job/company known)

    Dear Hiring Manager,

    [body paragraph 1]

    [body paragraph 2]

    ...

    Sincerely,

    Lorenzo White

The tailoring prompt explicitly omits greeting + sign-off so we can render them
consistently across PDF/DOCX without duplicating.
"""
from __future__ import annotations

import io
import logging
from datetime import datetime

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def build_docx(paragraphs: list[str], profile: dict, job: dict | None = None) -> bytes:
    """Render a cover letter as a DOCX with header + body + sign-off."""
    from docx import Document
    from docx.shared import Pt, Inches

    name = (profile.get("full_name") or "").strip() or "Cover Letter"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Header: name (bold), then contact lines ---
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(13)
    for line in _contact_lines(profile):
        doc.add_paragraph(line)

    doc.add_paragraph()  # spacer

    # --- Date ---
    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph()

    # --- Recipient (if we know the company) ---
    company = (job or {}).get("company_name") or (job or {}).get("company") or ""
    if company:
        doc.add_paragraph(f"Hiring Team, {company}")
        doc.add_paragraph()

    # --- Greeting ---
    doc.add_paragraph("Dear Hiring Manager,")
    doc.add_paragraph()

    # --- Body paragraphs ---
    for para in paragraphs:
        if para and para.strip():
            doc.add_paragraph(para.strip())
            doc.add_paragraph()

    # --- Sign-off ---
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph(name)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF (fpdf2 — same fallback the resume builder uses)
# ---------------------------------------------------------------------------

def build_pdf(paragraphs: list[str], profile: dict, job: dict | None = None) -> bytes:
    """Render a cover letter as a PDF."""
    from fpdf import FPDF

    name = (profile.get("full_name") or "").strip() or "Cover Letter"

    pdf = FPDF(format="Letter", unit="mm")
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=20)

    # Name (bold, slightly larger)
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 7, _safe(name), new_x="LMARGIN", new_y="NEXT")

    # Contact lines
    pdf.set_font("Helvetica", size=10)
    for line in _contact_lines(profile):
        pdf.cell(0, 5, _safe(line), new_x="LMARGIN", new_y="NEXT")

    pdf.ln(5)

    # Date
    pdf.set_font("Helvetica", size=11)
    pdf.cell(0, 5, datetime.now().strftime("%B %d, %Y"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Recipient
    company = (job or {}).get("company_name") or (job or {}).get("company") or ""
    if company:
        pdf.cell(0, 5, _safe(f"Hiring Team, {company}"), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    # Greeting
    pdf.cell(0, 5, "Dear Hiring Manager,", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    # Body paragraphs
    for para in paragraphs:
        if not para or not para.strip():
            continue
        pdf.multi_cell(0, 5, _safe(para.strip()))
        pdf.ln(3)

    # Sign-off
    pdf.cell(0, 5, "Sincerely,", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    pdf.cell(0, 5, _safe(name), new_x="LMARGIN", new_y="NEXT")

    out = pdf.output()
    return bytes(out) if isinstance(out, (bytes, bytearray)) else out.encode("latin-1", errors="replace")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _contact_lines(profile: dict) -> list[str]:
    """Pull the non-empty contact lines we want to display under the name."""
    lines: list[str] = []
    loc = (profile.get("location") or "").strip().rstrip(",").strip()
    phone = (profile.get("phone") or "").strip()
    email = (
        profile.get("resume_email")
        or profile.get("email")
        or profile.get("contact_email")
        or ""
    ).strip()
    linkedin = (profile.get("linkedin_url") or "").strip()

    if loc:
        lines.append(loc)
    if phone:
        lines.append(phone)
    if email:
        lines.append(email)
    if linkedin:
        # Strip protocol + trailing slash so it reads as "linkedin.com/in/..."
        clean = linkedin.replace("https://", "").replace("http://", "").rstrip("/")
        lines.append(clean)
    return lines


def _safe(text: str) -> str:
    """Replace Unicode punctuation that fpdf2's built-in Latin-1 fonts can't render."""
    return (
        (text or "")
        .replace("–", "-").replace("—", "-")
        .replace("‘", "'").replace("’", "'")
        .replace("“", '"').replace("”", '"')
        .replace("•", "-")
        .replace("…", "...")
        .encode("latin-1", errors="replace")
        .decode("latin-1")
    )
